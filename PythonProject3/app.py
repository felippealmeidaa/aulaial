from flask import render_template, request, jsonify, session, redirect, url_for, flash, send_file
from flask.app import Flask as FlaskApp
from typing import Any, cast
import google.generativeai as genai
import os
from datetime import datetime, timedelta
import sqlite3
from werkzeug.security import generate_password_hash, check_password_hash
import time
from dotenv import load_dotenv
import random
import json
from flask_mail import Mail, Message
from itsdangerous import URLSafeTimedSerializer, SignatureExpired, BadSignature
import threading
import re
import bleach
from functools import lru_cache
import io
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen.canvas import Canvas
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle

# ============================================
# 🆕 V5.1: IMPORTAR SCRAPER AVA COM NOVAS FUNÇÕES
# ============================================
try:
    from scraper_ava import (
        sincronizar_dados_ava,
        obter_ultima_sincronizacao,
        usuario_tem_cache
    )

    SCRAPER_DISPONIVEL = True
    print("✅ Scraper V5.1 carregado com sucesso!")
except ImportError as e:
    print(f"⚠️ AVISO: Scraper não disponível. Erro: {e}")
    SCRAPER_DISPONIVEL = False


    def sincronizar_dados_ava(user_id: int, *args, **kwargs) -> None:
        print("Scraper não disponível.")
        return None


    def obter_ultima_sincronizacao(user_id: int) -> datetime | None:
        return None


    def usuario_tem_cache(user_id: int) -> bool:
        return False

# ============================================
# 🆕 IMPORTAR SCRAPER LYCEUM V5.1
# ============================================
try:
    from scraper_lyceum import (
        sincronizar_dados_lyceum,
        sincronizar_dados_lyceum_v2,
        obter_ultima_sincronizacao_lyceum,
        usuario_tem_cache_lyceum
    )

    LYCEUM_DISPONIVEL = True
    print("✅ Scraper Lyceum V5.1 carregado com sucesso!")
except ImportError as e:
    print(f"⚠️ AVISO: Scraper Lyceum não disponível. Erro: {e}")
    LYCEUM_DISPONIVEL = False


    def sincronizar_dados_lyceum(user_id: int, *args, **kwargs) -> None:
        print("Scraper Lyceum não disponível.")
        return None

    def sincronizar_dados_lyceum_v2(user_id: int, *args, **kwargs) -> None:
        print("Scraper Lyceum V2 não disponível.")
        return None


    def obter_ultima_sincronizacao_lyceum(user_id: int) -> datetime | None:
        return None


    def usuario_tem_cache_lyceum(user_id: int) -> bool:
        return False

# ============================================
# CARREGAR VARIÁVEIS DE AMBIENTE
# ============================================
load_dotenv()
app = FlaskApp(__name__)

# ============================================
# 🆕 V5.1: CONTROLE DE SINCRONIZAÇÕES AVA
# ============================================
status_sincronizacao = {}
sincronizacoes_em_andamento = {}

# ============================================
# 🆕 V5.1: CONTROLE DE SINCRONIZAÇÕES LYCEUM
# ============================================
status_sincronizacao_lyceum = {}
sincronizacoes_em_andamento_lyceum = {}

# ============================================
# CONFIGURAÇÕES DE SEGURANÇA
# ============================================
app.secret_key = os.getenv('SECRET_KEY', 'chave_desenvolvimento_segura')
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')
DATABASE = os.getenv('DATABASE', 'unievangelica.db')

# Configurações do Flask-Mail
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USE_SSL'] = False
app.config['MAIL_USERNAME'] = os.getenv('MAIL_USERNAME')
app.config['MAIL_PASSWORD'] = os.getenv('MAIL_PASSWORD')
app.config['MAIL_DEFAULT_SENDER'] = os.getenv('MAIL_USERNAME')

mail = Mail()
mail.init_app(cast(Any, app))
s = URLSafeTimedSerializer(app.secret_key or 'chave_desenvolvimento_segura')

if not GEMINI_API_KEY:
    print("⚠️ AVISO: GEMINI_API_KEY não configurada no arquivo .env")

# ============================================
# CONFIGURAR GEMINI API
# ============================================
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)
    model = genai.GenerativeModel('models/gemini-2.5-flash')
    generation_config = {
        'temperature': 0.7,
        'top_p': 0.95,
        'top_k': 40,
        'max_output_tokens': 8000,
    }
else:
    model = None

# ============================================
# CONTROLE DE RATE LIMITING
# ============================================
last_request_time = {}
request_count = {}


def can_make_request(user_id, min_seconds=4, max_requests_per_minute=15):
    current_time = time.time()
    if user_id not in last_request_time:
        last_request_time[user_id] = current_time
        request_count[user_id] = {'count': 1, 'window_start': current_time}
        return True

    time_diff = current_time - last_request_time[user_id]
    if time_diff < min_seconds:
        return False

    if user_id in request_count:
        window_elapsed = current_time - request_count[user_id]['window_start']
        if window_elapsed > 60:
            request_count[user_id] = {'count': 1, 'window_start': current_time}
        else:
            if request_count[user_id]['count'] >= max_requests_per_minute:
                return False
            request_count[user_id]['count'] += 1

    last_request_time[user_id] = current_time
    return True


def get_wait_time(user_id):
    if user_id not in last_request_time:
        return 0
    current_time = time.time()
    time_diff = current_time - last_request_time[user_id]
    wait_time = max(0, 4 - time_diff)
    return int(wait_time)


# ============================================
# SANITIZAÇÃO DE HTML
# ============================================
def sanitizar_html(texto):
    """Remove tags HTML perigosas mantendo formatação segura"""
    if not texto: return ""
    tags_permitidas = ['b', 'i', 'u', 'p', 'br', 'strong', 'em', 'code', 'pre']
    return bleach.clean(texto, tags=frozenset(tags_permitidas), strip=True)


# ============================================
# BANCO DE DADOS
# ============================================
def get_db_connection():
    """Retorna uma conexão com o banco de dados."""
    conn = sqlite3.connect(DATABASE)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    with get_db_connection() as conn:
        c = conn.cursor()

        c.execute(''' 
            CREATE TABLE IF NOT EXISTS usuarios ( 
                id INTEGER PRIMARY KEY AUTOINCREMENT, 
                nome TEXT NOT NULL, 
                matricula TEXT UNIQUE NOT NULL, 
                cpf TEXT,  
                email TEXT UNIQUE NOT NULL, 
                curso TEXT NOT NULL, 
                senha TEXT NOT NULL, 
                dark_mode INTEGER DEFAULT 0,
                data_cadastro TIMESTAMP DEFAULT CURRENT_TIMESTAMP 
            ) 
        ''')

        c.execute(''' 
            CREATE TABLE IF NOT EXISTS conteudos_ava ( 
                id INTEGER PRIMARY KEY AUTOINCREMENT,  
                usuario_id INTEGER,  
                disciplina TEXT,  
                conteudo_texto TEXT,  
                data_extracao TIMESTAMP DEFAULT CURRENT_TIMESTAMP 
            ) 
        ''')

        c.execute(''' 
            CREATE TABLE IF NOT EXISTS posts ( 
                id INTEGER PRIMARY KEY AUTOINCREMENT, 
                post_id TEXT UNIQUE NOT NULL, 
                curso TEXT NOT NULL, 
                tipo TEXT NOT NULL, 
                titulo TEXT NOT NULL, 
                conteudo TEXT NOT NULL,
                tags TEXT,
                usuario_id INTEGER NOT NULL, 
                nome_usuario TEXT NOT NULL, 
                data_criacao TIMESTAMP DEFAULT CURRENT_TIMESTAMP, 
                FOREIGN KEY (usuario_id) REFERENCES usuarios (id) 
            ) 
        ''')

        c.execute(''' 
            CREATE TABLE IF NOT EXISTS curtidas ( 
                id INTEGER PRIMARY KEY AUTOINCREMENT, 
                post_id TEXT NOT NULL, 
                usuario_id INTEGER NOT NULL, 
                data_curtida TIMESTAMP DEFAULT CURRENT_TIMESTAMP, 
                FOREIGN KEY (usuario_id) REFERENCES usuarios (id), 
                UNIQUE(post_id, usuario_id) 
            ) 
        ''')

        c.execute(''' 
            CREATE TABLE IF NOT EXISTS comentarios ( 
                id INTEGER PRIMARY KEY AUTOINCREMENT, 
                post_id TEXT NOT NULL, 
                usuario_id INTEGER NOT NULL, 
                nome_usuario TEXT NOT NULL, 
                comentario TEXT NOT NULL, 
                data_comentario TIMESTAMP DEFAULT CURRENT_TIMESTAMP, 
                FOREIGN KEY (usuario_id) REFERENCES usuarios (id) 
            ) 
        ''')

        c.execute(''' 
            CREATE TABLE IF NOT EXISTS historico_chat ( 
                id INTEGER PRIMARY KEY AUTOINCREMENT, 
                usuario_id INTEGER NOT NULL, 
                mensagem TEXT NOT NULL, 
                resposta TEXT NOT NULL, 
                data_hora TIMESTAMP DEFAULT CURRENT_TIMESTAMP, 
                FOREIGN KEY (usuario_id) REFERENCES usuarios (id) 
            ) 
        ''')

        c.execute(''' 
            CREATE TABLE IF NOT EXISTS eventos_calendario ( 
                id INTEGER PRIMARY KEY AUTOINCREMENT, 
                usuario_id INTEGER NOT NULL, 
                titulo TEXT NOT NULL, 
                descricao TEXT, 
                data_evento DATE NOT NULL, 
                hora_evento TIME, 
                tipo TEXT NOT NULL, 
                cor TEXT DEFAULT '#4a90e2', 
                alerta INTEGER DEFAULT 0, 
                minutos_antes_alerta INTEGER DEFAULT 30, 
                data_criacao TIMESTAMP DEFAULT CURRENT_TIMESTAMP, 
                FOREIGN KEY (usuario_id) REFERENCES usuarios (id) 
            ) 
        ''')

        c.execute(''' 
            CREATE TABLE IF NOT EXISTS notas_aluno ( 
                id INTEGER PRIMARY KEY AUTOINCREMENT, 
                usuario_id INTEGER NOT NULL, 
                disciplina TEXT NOT NULL, 
                va1 REAL DEFAULT 0, 
                va2 REAL DEFAULT 0, 
                va3 REAL DEFAULT 0, 
                media REAL DEFAULT 0, 
                situacao TEXT DEFAULT 'Cursando', 
                FOREIGN KEY (usuario_id) REFERENCES usuarios (id) 
            ) 
        ''')

        c.execute(''' 
            CREATE TABLE IF NOT EXISTS faltas_aluno ( 
                id INTEGER PRIMARY KEY AUTOINCREMENT, 
                usuario_id INTEGER NOT NULL, 
                disciplina TEXT NOT NULL, 
                total_faltas INTEGER DEFAULT 0, 
                total_aulas INTEGER DEFAULT 60, 
                percentual_presenca REAL DEFAULT 100, 
                FOREIGN KEY (usuario_id) REFERENCES usuarios (id) 
            ) 
        ''')

        c.execute(''' 
            CREATE TABLE IF NOT EXISTS notificacoes ( 
                id INTEGER PRIMARY KEY AUTOINCREMENT, 
                usuario_id INTEGER NOT NULL, 
                tipo TEXT NOT NULL, 
                mensagem TEXT NOT NULL, 
                link TEXT, 
                lida INTEGER DEFAULT 0, 
                data_criacao TIMESTAMP DEFAULT CURRENT_TIMESTAMP, 
                FOREIGN KEY (usuario_id) REFERENCES usuarios (id) 
            ) 
        ''')

        try:
            c.execute("SELECT tags FROM posts LIMIT 1")
        except sqlite3.OperationalError:
            c.execute("ALTER TABLE posts ADD COLUMN tags TEXT")

        try:
            c.execute("SELECT dark_mode FROM usuarios LIMIT 1")
        except sqlite3.OperationalError:
            c.execute("ALTER TABLE usuarios ADD COLUMN dark_mode INTEGER DEFAULT 0")

        try:
            c.execute("SELECT ultima_atualizacao_lyceum FROM usuarios LIMIT 1")
        except sqlite3.OperationalError:
            c.execute("ALTER TABLE usuarios ADD COLUMN ultima_atualizacao_lyceum TEXT")

        try:
            c.execute("SELECT senha_lyceum FROM usuarios LIMIT 1")
        except sqlite3.OperationalError:
            c.execute("ALTER TABLE usuarios ADD COLUMN senha_lyceum TEXT")

        conn.commit()
    print("✅ Banco de dados inicializado/verificado com sucesso!")


# ============================================
# SISTEMA DE NOTIFICAÇÕES
# ============================================
def criar_notificacao(usuario_id, tipo, mensagem, link=None):
    """Cria uma notificação para o usuário"""
    try:
        with get_db_connection() as conn:
            conn.execute('''
                INSERT INTO notificacoes (usuario_id, tipo, mensagem, link, lida)
                VALUES (?, ?, ?, ?, 0)
            ''', (usuario_id, tipo, mensagem, link))
            conn.commit()
    except Exception as e:
        print(f"[ERROR] Erro ao criar notificacao: {e}")


@app.route('/api/notificacoes')
def buscar_notificacoes():
    """Busca notificações do usuário"""
    if 'user_id' not in session:
        return jsonify({'error': 'Não autorizado'}), 401

    user_id = session['user_id']

    try:
        with get_db_connection() as conn:
            c = conn.cursor()
            c.execute('''
                SELECT id, tipo, mensagem, link, lida, 
                       strftime('%d/%m/%Y %H:%M', data_criacao) as data_formatada
                FROM notificacoes
                WHERE usuario_id = ?
                ORDER BY data_criacao DESC
                LIMIT 20
            ''', (user_id,))

            notificacoes = []
            for row in c.fetchall():
                notificacoes.append({
                    'id': row['id'],
                    'tipo': row['tipo'],
                    'mensagem': row['mensagem'],
                    'link': row['link'],
                    'lida': row['lida'],
                    'data': row['data_formatada']
                })

            c.execute('SELECT COUNT(*) as total FROM notificacoes WHERE usuario_id = ? AND lida = 0', (user_id,))
            result = c.fetchone()
            nao_lidas = result['total'] if result else 0

        return jsonify({
            'success': True,
            'notificacoes': notificacoes,
            'nao_lidas': nao_lidas
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/notificacoes/marcar_lida', methods=['POST'])
def marcar_notificacao_lida():
    """Marca notificação como lida"""
    if 'user_id' not in session:
        return jsonify({'error': 'Não autorizado'}), 401

    data = request.get_json(silent=True) or {}
    notif_id = data.get('notificacao_id')

    try:
        with get_db_connection() as conn:
            conn.execute('UPDATE notificacoes SET lida = 1 WHERE id = ? AND usuario_id = ?',
                         (notif_id, session['user_id']))
            conn.commit()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/notificacoes/marcar_todas_lidas', methods=['POST'])
def marcar_todas_lidas():
    """Marca todas as notificações como lidas"""
    if 'user_id' not in session:
        return jsonify({'error': 'Não autorizado'}), 401

    try:
        with get_db_connection() as conn:
            conn.execute('UPDATE notificacoes SET lida = 1 WHERE usuario_id = ?', (session['user_id'],))
            conn.commit()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ============================================
# CACHE DE DADOS
# ============================================
@lru_cache(maxsize=128)
def get_eventos_cache_key(user_id, timestamp):
    """Cache de eventos com invalidação por timestamp"""
    with get_db_connection() as conn:
        c = conn.cursor()
        c.execute('''
            SELECT id, titulo, descricao, data_evento, hora_evento, tipo, cor, alerta, minutos_antes_alerta
            FROM eventos_calendario
            WHERE usuario_id = ?
            ORDER BY data_evento ASC
        ''', (user_id,))
        return c.fetchall()


def limpar_cache_eventos():
    """Limpa o cache de eventos"""
    get_eventos_cache_key.cache_clear()


# ============================================
# DADOS ACADÊMICOS
# ============================================
DADOS_ACADEMICOS = """ 
INFORMAÇÕES DA UNIEVANGELICA - CURSO DE INTELIGÊNCIA ARTIFICIAL 
Data de Atualização: 2025 

HORÁRIOS DAS AULAS - SEMESTRE 2025.2 
SEGUNDA-FEIRA: 
• 19:00 - 20:40: Cidadania ética e espiritualidade 
• 21:00 - 22:40: Fundamentos matemáticos para computação 

TERÇA-FEIRA: 
• 19:00 - 22:40: Introdução a engenharia de soluções 

QUARTA-FEIRA: 
• 19:00 - 22:40: Fundamentos matemáticos para computação 

QUINTA-FEIRA: 
• 19:00 - 22:40: Fundamentos de computação e infraestrutura 

SEXTA-FEIRA: 
• 19:00 - 22:40: Fundamento de engenharia de dados 

EAD: 
• Introdução a língua portuguesa (modalidade online) 

DISCIPLINAS DO CURSO: 
1. Cidadania ética e espiritualidade 
2. Introdução a engenharia de soluções 
3. Fundamentos matemáticos para computação 
4. Fundamentos de computação e infraestrutura 
5. Fundamentos de engenharia de dados 
6. Leitura e interpretação de texto 

CORPO DOCENTE: 
• Coordenadora: Natasha Sophie Campos 
• Prof. Holehon Santos Campos - Cidadania ética 
• Prof. Eder José Almeida da Silva - Eng. de soluções 
• Prof. Henrique Valle de Lima - Fundamentos matemáticos 
• Prof. Jeferson Silva Araújo - Computação e infraestrutura 
• Prof. Fábio Pereira Botelho - Engenharia de dados 
• Prof. Leonardo Rodrigues de Souza - Leitura e interpretação 

CALENDÁRIO DE PROVAS 2025.2: 
• 1ª VA: 15 a 20 de Setembro 
• 2ª VA: 27 de Outubro a 01 de Novembro 
• 3ª VA: 08 a 13 de Dezembro 
• Substitutivas 1ª e 2ª VA: 11 e 12 de Novembro 
• Substitutivas 3ª VA: 16 e 17 de Dezembro 
"""

HORARIOS_AULAS = {
    'Segunda-feira': [
        {'horario': '19:00 - 20:40', 'disciplina': 'Cidadania ética e espiritualidade',
         'professor': 'Holehon Santos Campos'},
        {'horario': '21:00 - 22:40', 'disciplina': 'Fundamentos matemáticos para computação',
         'professor': 'Henrique Valle de Lima'}
    ],
    'Terça-feira': [
        {'horario': '19:00 - 22:40', 'disciplina': 'Introdução a engenharia de soluções',
         'professor': 'Eder José Almeida da Silva'}
    ],
    'Quarta-feira': [
        {'horario': '19:00 - 22:40', 'disciplina': 'Fundamentos matemáticos para computação',
         'professor': 'Henrique Valle de Lima'}
    ],
    'Quinta-feira': [
        {'horario': '19:00 - 22:40', 'disciplina': 'Fundamentos de computação e infraestrutura',
         'professor': 'Jeferson Silva Araújo'}
    ],
    'Sexta-feira': [
        {'horario': '19:00 - 22:40', 'disciplina': 'Fundamento de engenharia de dados',
         'professor': 'Fábio Pereira Botelho'}
    ],
    'EAD': [
        {'horario': 'Online', 'disciplina': 'Leitura e interpretação de texto',
         'professor': 'Leonardo Rodrigues de Souza'}
    ]
}

EVENTOS_ACADEMICOS = [
    {'title': '🎉 Feriado Municipal', 'start': '2025-07-26', 'color': '#e63946', 'tipo': 'feriado', 'allDay': True},
    {'title': '🎓 Colação de Grau Unificada', 'start': '2025-07-30', 'color': '#4a90e2', 'tipo': 'evento',
     'allDay': True},
    {'title': '🎂 Aniversário de Anápolis', 'start': '2025-07-31', 'color': '#e63946', 'tipo': 'feriado',
     'allDay': True},
    {'title': '📚 INÍCIO DAS AULAS', 'start': '2025-08-04', 'color': '#28a745', 'tipo': 'importante', 'allDay': True},
    {'title': '🎓 Dia do Estudante', 'start': '2025-08-11', 'color': '#ffc107', 'tipo': 'comemorativo', 'allDay': True},
    {'title': '🏆 Prêmio Mérito Científico', 'start': '2025-08-19', 'color': '#4a90e2', 'tipo': 'evento',
     'allDay': True},
    {'title': '⚠️ Inclusão/Exclusão Disciplinas', 'start': '2025-08-29', 'color': '#ff6b6b', 'tipo': 'prazo',
     'allDay': True},
    {'title': '🇧🇷 Independência do Brasil', 'start': '2025-09-07', 'color': '#e63946', 'tipo': 'feriado',
     'allDay': True},
    {'title': '📝 INÍCIO 1ª VA', 'start': '2025-09-15', 'color': '#dc3545', 'tipo': 'prova', 'allDay': True},
    {'title': '📝 1ª VA', 'start': '2025-09-16', 'color': '#dc3545', 'tipo': 'prova', 'allDay': True},
    {'title': '📝 1ª VA', 'start': '2025-09-17', 'color': '#dc3545', 'tipo': 'prova', 'allDay': True},
    {'title': '📝 1ª VA', 'start': '2025-09-18', 'color': '#dc3545', 'tipo': 'prova', 'allDay': True},
    {'title': '📝 1ª VA', 'start': '2025-09-19', 'color': '#dc3545', 'tipo': 'prova', 'allDay': True},
    {'title': '📝 FIM 1ª VA', 'start': '2025-09-20', 'color': '#dc3545', 'tipo': 'prova', 'allDay': True},
    {'title': '⏰ Lançamento notas 1ª VA', 'start': '2025-10-03', 'color': '#ff6b6b', 'tipo': 'prazo', 'allDay': True},
    {'title': '🇧🇷 Nossa Senhora Aparecida', 'start': '2025-10-12', 'color': '#e63946', 'tipo': 'feriado',
     'allDay': True},
    {'title': '👨‍🏫 Dia do Professor', 'start': '2025-10-15', 'color': '#ffc107', 'tipo': 'comemorativo',
     'allDay': True},
    {'title': '📝 INÍCIO 2ª VA', 'start': '2025-10-27', 'color': '#dc3545', 'tipo': 'prova', 'allDay': True},
    {'title': '📝 2ª VA', 'start': '2025-10-28', 'color': '#dc3545', 'tipo': 'prova', 'allDay': True},
    {'title': '📝 2ª VA', 'start': '2025-10-29', 'color': '#dc3545', 'tipo': 'prova', 'allDay': True},
    {'title': '📝 2ª VA', 'start': '2025-10-30', 'color': '#dc3545', 'tipo': 'prova', 'allDay': True},
    {'title': '⚠️ Trancamento de Matrícula', 'start': '2025-10-31', 'color': '#ff6b6b', 'tipo': 'prazo',
     'allDay': True},
    {'title': '📝 FIM 2ª VA', 'start': '2025-11-01', 'color': '#dc3545', 'tipo': 'prova', 'allDay': True},
    {'title': '🕯️ Finados', 'start': '2025-11-02', 'color': '#e63946', 'tipo': 'feriado', 'allDay': True},
    {'title': '🔄 Substitutivas 1ª e 2ª VA', 'start': '2025-11-11', 'color': '#dc3545', 'tipo': 'prova', 'allDay': True},
    {'title': '🔄 Substitutivas 1ª e 2ª VA', 'start': '2025-11-12', 'color': '#dc3545', 'tipo': 'prova', 'allDay': True},
    {'title': '🇧🇷 Proclamação da República', 'start': '2025-11-15', 'color': '#e63946', 'tipo': 'feriado',
     'allDay': True},
    {'title': '✊ Consciência Negra', 'start': '2025-11-20', 'color': '#e63946', 'tipo': 'feriado', 'allDay': True},
    {'title': '📝 INÍCIO 3ª VA', 'start': '2025-12-08', 'color': '#dc3545', 'tipo': 'prova', 'allDay': True},
    {'title': '📝 3ª VA', 'start': '2025-12-09', 'color': '#dc3545', 'tipo': 'prova', 'allDay': True},
    {'title': '📝 3ª VA', 'start': '2025-12-10', 'color': '#dc3545', 'tipo': 'prova', 'allDay': True},
    {'title': '📝 3ª VA', 'start': '2025-12-11', 'color': '#dc3545', 'tipo': 'prova', 'allDay': True},
    {'title': '📝 3ª VA', 'start': '2025-12-12', 'color': '#dc3545', 'tipo': 'prova', 'allDay': True},
    {'title': '📝 FIM 3ª VA', 'start': '2025-12-13', 'color': '#dc3545', 'tipo': 'prova', 'allDay': True},
    {'title': '🔄 Substitutivas 3ª VA', 'start': '2025-12-16', 'color': '#dc3545', 'tipo': 'prova', 'allDay': True},
    {'title': '🔄 Substitutivas 3ª VA', 'start': '2025-12-17', 'color': '#dc3545', 'tipo': 'prova', 'allDay': True},
    {'title': '🎄 Natal', 'start': '2025-12-25', 'color': '#e63946', 'tipo': 'feriado', 'allDay': True},
    {'title': '🎆 Ano Novo', 'start': '2026-01-01', 'color': '#e63946', 'tipo': 'feriado', 'allDay': True},
]


def executar_sincronizacao_monitorada(user_id, matricula, cpf, forcar=False):
    """V5.1: Sincronização AVA em thread"""
    try:
        print(f"[SYNC AVA] Iniciando sync para ID {user_id}")
        status_sincronizacao[user_id] = 'em_andamento'
        sincronizacoes_em_andamento[user_id] = True

        sincronizar_dados_ava(user_id, matricula, cpf, forcar_atualizacao=forcar)

        print(f"[SYNC AVA] Sync finalizado para ID {user_id}")
        status_sincronizacao[user_id] = 'concluido'
        sincronizacoes_em_andamento[user_id] = False
    except Exception as e:
        print(f"[ERROR AVA] Erro no sync: {e}")
        status_sincronizacao[user_id] = 'erro'
        sincronizacoes_em_andamento[user_id] = False


def executar_sincronizacao_monitorada_lyceum(user_id, matricula, senha_lyceum, forcar=False):
    """V5.1: Sincronização Lyceum em thread"""
    try:
        print(f"[SYNC LYCEUM] Iniciando sync para ID {user_id}")
        status_sincronizacao_lyceum[user_id] = 'em_andamento'
        sincronizacoes_em_andamento_lyceum[user_id] = True

        sincronizar_dados_lyceum_v2(user_id, matricula, senha_lyceum, forcar_atualizacao=forcar)

        print(f"[SYNC LYCEUM] Sync finalizado para ID {user_id}")
        status_sincronizacao_lyceum[user_id] = 'concluido'
        sincronizacoes_em_andamento_lyceum[user_id] = False
    except Exception as e:
        print(f"[ERROR LYCEUM] Erro no sync: {e}")
        status_sincronizacao_lyceum[user_id] = 'erro'
        sincronizacoes_em_andamento_lyceum[user_id] = False


def gerar_notas_ficticias(usuario_id):
    """Gera notas fictícias para o aluno"""
    disciplinas = [
        'Cidadania ética e espiritualidade',
        'Introdução a engenharia de soluções',
        'Fundamentos matemáticos para computação',
        'Fundamentos de computação e infraestrutura',
        'Fundamentos de engenharia de dados',
        'Leitura e interpretação de texto'
    ]
    with get_db_connection() as conn:
        c = conn.cursor()
        for disciplina in disciplinas:
            c.execute('SELECT * FROM notas_aluno WHERE usuario_id = ? AND disciplina = ?', (usuario_id, disciplina))
            if not c.fetchone():
                va1 = round(random.uniform(5.0, 10.0), 1)
                va2 = round(random.uniform(5.0, 10.0), 1)
                va3 = round(random.uniform(5.0, 10.0), 1)
                media = round((va1 + va2 + va3) / 3, 1)
                situacao = 'Aprovado' if media >= 6.0 else 'Reprovado' if media > 0 else 'Cursando'
                c.execute(''' 
                    INSERT INTO notas_aluno (usuario_id, disciplina, va1, va2, va3, media, situacao) 
                    VALUES (?, ?, ?, ?, ?, ?, ?) 
                ''', (usuario_id, disciplina, va1, va2, va3, media, situacao))
        conn.commit()


def gerar_faltas_ficticias(usuario_id):
    """Gera faltas fictícias para o aluno"""
    disciplinas = [
        'Cidadania ética e espiritualidade',
        'Introdução a engenharia de soluções',
        'Fundamentos matemáticos para computação',
        'Fundamentos de computação e infraestrutura',
        'Fundamentos de engenharia de dados',
        'Leitura e interpretação de texto'
    ]
    with get_db_connection() as conn:
        c = conn.cursor()
        for disciplina in disciplinas:
            c.execute('SELECT * FROM faltas_aluno WHERE usuario_id = ? AND disciplina = ?', (usuario_id, disciplina))
            if not c.fetchone():
                total_faltas = random.randint(0, 10)
                total_aulas = 60
                percentual_presenca = round(((total_aulas - total_faltas) / total_aulas) * 100, 1)
                c.execute(''' 
                    INSERT INTO faltas_aluno (usuario_id, disciplina, total_faltas, total_aulas, percentual_presenca) 
                    VALUES (?, ?, ?, ?, ?) 
                ''', (usuario_id, disciplina, total_faltas, total_aulas, percentual_presenca))
        conn.commit()


def salvar_historico_chat(usuario_id, mensagem, resposta):
    try:
        with get_db_connection() as conn:
            c = conn.cursor()
            c.execute(''' 
                INSERT INTO historico_chat (usuario_id, mensagem, resposta) 
                VALUES (?, ?, ?) 
            ''', (usuario_id, mensagem, resposta))
            conn.commit()
    except Exception as e:
        print(f"[ERROR] Erro ao salvar historico: {e}")


def montar_contexto_comunidade(limite_posts=6):
    try:
        with get_db_connection() as conn:
            c = conn.cursor()
            c.execute(''' 
                SELECT post_id, curso, titulo, conteudo, nome_usuario, 
                       strftime('%d/%m/%Y %H:%M', data_criacao) as data_formatada 
                FROM posts 
                WHERE tipo = 'duvida' 
                ORDER BY data_criacao DESC 
                LIMIT ? 
            ''', (limite_posts,))
            posts = c.fetchall()

            if not posts:
                return "Nenhuma dúvida registrada na comunidade."

            bloco = []
            for post in posts:
                post_id = post['post_id']
                curso = post['curso']
                curso_label = {
                    'ia': 'Inteligência Artificial',
                    'ads': 'Análise e Desenvolvimento de Sistemas',
                    'es': 'Engenharia de Software'
                }.get(curso, curso)

                bloco.append(f"[DÚVIDA - {curso_label}] {post['titulo']}")
                bloco.append(f"Pergunta: {post['conteudo']}")

                c.execute(''' 
                    SELECT nome_usuario, comentario 
                    FROM comentarios 
                    WHERE post_id = ? 
                    ORDER BY data_comentario ASC 
                    LIMIT 5 
                ''', (post_id,))
                comentarios = c.fetchall()
                if comentarios:
                    bloco.append("Respostas:")
                    for com in comentarios:
                        bloco.append(f"- {com['nome_usuario']}: {com['comentario']}")
                bloco.append("")

            return "\n".join(bloco)
    except Exception as e:
        print(f"[ERROR] Erro ao montar contexto: {e}")
        return "Não foi possível carregar dúvidas."


def criar_evento_rapido_via_chat(usuario_id, mensagem):
    texto = mensagem.strip()
    gatilho = "criar evento:"

    if not texto.lower().startswith(gatilho):
        return None

    try:
        conteudo = texto[len(gatilho):].strip()
        partes = [p.strip() for p in conteudo.split('|')]

        if len(partes) < 2:
            return "Para criar evento: criar evento: Título | AAAA-MM-DD | HH:MM (opcional)."

        titulo = partes[0]
        data_evento = partes[1]
        hora_evento = partes[2] if len(partes) >= 3 else ''

        try:
            datetime.strptime(data_evento, "%Y-%m-%d")
        except ValueError:
            return "Data inválida. Use: AAAA-MM-DD."

        with get_db_connection() as conn:
            conn.execute(''' 
                INSERT INTO eventos_calendario ( 
                    usuario_id, titulo, descricao, data_evento, hora_evento, 
                    tipo, cor, alerta, minutos_antes_alerta 
                ) 
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?) 
            ''', (usuario_id, titulo, 'Criado via chat', data_evento, hora_evento, 'pessoal', '#4a90e2', 0, 30))
            conn.commit()

        limpar_cache_eventos()

        if hora_evento:
            return f"[OK] Evento criado: **{titulo}** em **{data_evento} as {hora_evento}**."
        else:
            return f"[OK] Evento criado: **{titulo}** em **{data_evento}**."
    except Exception as e:
        print(f"[ERROR] Erro ao criar evento: {e}")
        return "[ERROR] Erro ao criar evento."


# ============================================
# ROTAS
# ============================================
@app.route('/')
def index():
    if 'user_id' in session:
        return redirect(url_for('dashboard'))
    erro = session.pop('login_erro', None)
    sucesso = session.pop('login_sucesso', None)
    return render_template('index.html', erro=erro, sucesso=sucesso)


@app.route('/cadastro')
def cadastro():
    return render_template('cadastro.html')


@app.route('/cadastro', methods=['POST'])
def cadastrar():
    nome = request.form['nome']
    matricula = request.form['matricula']
    cpf = request.form['cpf']
    email = request.form['email']
    curso = request.form['curso']
    senha = request.form['password']
    confirm_senha = request.form['confirm_password']

    if senha != confirm_senha:
        return render_template('cadastro.html', erro='As senhas não coincidem!')

    try:
        with get_db_connection() as conn:
            c = conn.cursor()
            c.execute(''' 
                INSERT INTO usuarios (nome, matricula, cpf, email, curso, senha) 
                VALUES (?, ?, ?, ?, ?, ?) 
            ''', (nome, matricula, cpf, email, curso, generate_password_hash(senha)))
            conn.commit()
            usuario_id = c.lastrowid

        gerar_notas_ficticias(usuario_id)
        gerar_faltas_ficticias(usuario_id)

        session['login_sucesso'] = 'Cadastro realizado! Faça login.'
        return redirect(url_for('index'))
    except sqlite3.IntegrityError:
        return render_template('cadastro.html', erro='Matrícula ou e-mail já cadastrados!')


# ============================================
# 🆕 V5.1 CORRETO: ROTA DE LOGIN
# ============================================
@app.route('/login', methods=['POST'])
def login():
    """
    🆕 V5.1 CORRETO:
    - 1º LOGIN (usuário novo sem dados): Scraping AUTOMÁTICO + tela loading
    - 2º+ LOGIN (usuário com dados): Login RÁPIDO (< 5 seg)
    - BOTÃO: Para RE-SINCRONIZAR quando quiser
    """
    login_input = request.form.get('login') or request.form.get('matricula') or request.form.get('email')
    senha = request.form.get('password')

    if not login_input or not senha:
        session['login_erro'] = 'Preencha todos os campos!'
        return redirect(url_for('index'))

    with get_db_connection() as conn:
        c = conn.cursor()
        c.execute(''' 
            SELECT id, nome, matricula, email, curso, senha, cpf, dark_mode  
            FROM usuarios  
            WHERE email = ? OR matricula = ? 
        ''', (login_input, login_input))
        usuario = c.fetchone()

    if usuario and check_password_hash(usuario['senha'], senha):
        # ============================================
        # SALVA DADOS NA SESSÃO (CRÍTICO!)
        # ============================================
        session['user_id'] = usuario['id']
        session['user_nome'] = usuario['nome']
        session['user_curso'] = usuario['curso']
        session['dark_mode'] = usuario['dark_mode']
        session['matricula'] = usuario['matricula']
        session['cpf'] = usuario['cpf']

        print(f"[LOGIN] Login bem-sucedido: {usuario['nome']} (ID: {usuario['id']})")

        # ============================================
        # DECISÃO: SCRAPING AUTOMÁTICO OU LOGIN RÁPIDO?
        # ============================================
        if usuario['cpf'] and SCRAPER_DISPONIVEL:
            tem_dados = usuario_tem_cache(usuario['id'])

            if tem_dados:
                # ============================================
                # ✅ USUÁRIO COM DADOS: LOGIN RÁPIDO
                # ============================================
                ultima_sync = obter_ultima_sincronizacao(usuario['id'])
                if ultima_sync:
                    print(f"[LOGIN] Login rápido (Cache válido)")
                    print(f"[LOGIN] Última sincronização: {ultima_sync.strftime('%d/%m/%Y às %H:%M')}")
                else:
                    print(f"[LOGIN] Login rápido (Dados existem)")

                status_sincronizacao[usuario['id']] = 'concluido'
                gerar_notas_ficticias(usuario['id'])
                gerar_faltas_ficticias(usuario['id'])

                return redirect(url_for('dashboard'))

            else:
                status_sincronizacao[usuario['id']] = 'iniciando'
                status_sincronizacao_lyceum[usuario['id']] = 'iniciando'
                sincronizacoes_em_andamento[usuario['id']] = True
                sincronizacoes_em_andamento_lyceum[usuario['id']] = True

                thread_ava = threading.Thread(
                    target=executar_sincronizacao_monitorada,
                    args=(usuario['id'], usuario['matricula'], usuario['cpf'], True),
                    daemon=True
                )
                thread_ava.start()

                senha_lyceum = None
                try:
                    with get_db_connection() as conn:
                        c = conn.cursor()
                        c.execute('SELECT senha_lyceum, cpf FROM usuarios WHERE id = ?', (usuario['id'],))
                        row = c.fetchone()
                        if row:
                            senha_lyceum = row['senha_lyceum'] or row['cpf']
                except Exception:
                    senha_lyceum = usuario['cpf']

                if senha_lyceum:
                    thread_lyc = threading.Thread(
                        target=executar_sincronizacao_monitorada_lyceum,
                        args=(usuario['id'], usuario['matricula'], senha_lyceum, True),
                        daemon=True
                    )
                    thread_lyc.start()

                return redirect(url_for('dashboard'))

        gerar_notas_ficticias(usuario['id'])
        gerar_faltas_ficticias(usuario['id'])
        return redirect(url_for('dashboard'))

    else:
        session['login_erro'] = 'E-mail/matrícula ou senha incorretos!'
        return redirect(url_for('index'))


@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('index'))


@app.route('/esqueci-senha', methods=['GET', 'POST'])
def esqueci_senha():
    if request.method == 'POST':
        email = request.form.get('email')

        with get_db_connection() as conn:
            c = conn.cursor()
            c.execute('SELECT email FROM usuarios WHERE email = ?', (email,))
            usuario = c.fetchone()

        if usuario:
            token = s.dumps(email, salt='email-recuperacao')
            link = url_for('redefinir_senha', token=token, _external=True)

            try:
                msg = Message('Recuperação de Senha - IAUniev', recipients=[email] if email else None)
                msg.body = f'Clique no link: {link}\n\nExpira em 1 hora.'
                msg.html = f"""<p>Clique no botão:</p><a href="{link}" style="background:#0056b3; color:white; padding:10px 20px; text-decoration:none; border-radius:5px;">Redefinir Senha</a>"""
                mail.send(msg)
                flash('E-mail enviado!', 'success')
            except Exception as e:
                print(f"Erro: {e}")
                print(f"🔗 LINK (DEV): {link}")
                flash('Erro ao enviar. (Veja console)', 'warning')
        else:
            flash('E-mail não encontrado.', 'error')

    return render_template('esqueci_senha.html')


@app.route('/redefinir-senha/<token>', methods=['GET', 'POST'])
def redefinir_senha(token):
    try:
        email = s.loads(token, salt='email-recuperacao', max_age=3600)
    except SignatureExpired:
        return render_template('redefinir_senha.html', erro='Link expirou!', token=None)
    except BadSignature:
        return render_template('redefinir_senha.html', erro='Link inválido!', token=None)

    if request.method == 'POST':
        nova_senha = request.form.get('password')
        confirm_senha = request.form.get('confirm_password')

        if not nova_senha or not confirm_senha or nova_senha != confirm_senha:
            return render_template('redefinir_senha.html', erro='Senhas não coincidem!', token=token)

        senha_hash = generate_password_hash(nova_senha)

        with get_db_connection() as conn:
            conn.execute('UPDATE usuarios SET senha = ? WHERE email = ?', (senha_hash, email))
            conn.commit()

        session['login_sucesso'] = 'Senha redefinida!'
        return redirect(url_for('index'))

    return render_template('redefinir_senha.html', token=token)


@app.route('/dashboard')
def dashboard():
    if 'user_id' not in session:
        return redirect(url_for('index'))

    nome = session.get('user_nome', 'Usuário')
    curso = session.get('user_curso', 'Não especificado')
    user_id = session['user_id']
    dark_mode = session.get('dark_mode', 0)

    with get_db_connection() as conn:
        c = conn.cursor()
        posts_por_curso = {}
        cursos = ['ia', 'ads', 'es']
        for curso_id in cursos:
            c.execute(''' 
                SELECT post_id, tipo, titulo, conteudo, tags, nome_usuario,  
                       strftime('%d/%m/%Y %H:%M', data_criacao) as data_formatada, 
                       usuario_id 
                FROM posts 
                WHERE curso = ? 
                ORDER BY data_criacao DESC 
            ''', (curso_id,))
            posts = []
            for row in c.fetchall():
                posts.append({
                    'post_id': row['post_id'],
                    'tipo': row['tipo'],
                    'titulo': row['titulo'],
                    'conteudo': row['conteudo'],
                    'tags': row['tags'] or '',
                    'nome_usuario': row['nome_usuario'],
                    'data': row['data_formatada'],
                    'e_meu': row['usuario_id'] == user_id
                })
            posts_por_curso[curso_id] = posts

        c.execute(''' 
            SELECT disciplina, va1, va2, va3, media, situacao 
            FROM notas_aluno 
            WHERE usuario_id = ? 
            ORDER BY disciplina 
        ''', (user_id,))
        notas = []
        for row in c.fetchall():
            notas.append({
                'disciplina': row['disciplina'],
                'va1': row['va1'],
                'va2': row['va2'],
                'va3': row['va3'],
                'media': row['media'],
                'situacao': row['situacao']
            })

        c.execute(''' 
            SELECT disciplina, total_faltas, total_aulas, percentual_presenca 
            FROM faltas_aluno 
            WHERE usuario_id = ? 
            ORDER BY disciplina 
        ''', (user_id,))
        faltas = []
        for row in c.fetchall():
            faltas.append({
                'disciplina': row['disciplina'],
                'total_faltas': row['total_faltas'],
                'total_aulas': row['total_aulas'],
                'percentual_presenca': row['percentual_presenca']
            })

        # ============================================
        # 🆕 V7.0: BUSCAR HORÁRIOS DO SCRAPER
        # ============================================
        horarios_aulas_dinamicos = {}
        try:
            c.execute('''
                SELECT dia_semana, dia_nome, disciplina, horario_inicio, horario_fim, local, professor
                FROM horarios_aluno
                WHERE usuario_id = ?
                ORDER BY dia_semana, horario_inicio
            ''', (user_id,))
            
            horarios_rows = c.fetchall()
            
            if horarios_rows:
                # Organizar por dia da semana
                for row in horarios_rows:
                    dia_nome = row['dia_nome'] or f"Dia {row['dia_semana']}"
                    
                    if dia_nome not in horarios_aulas_dinamicos:
                        horarios_aulas_dinamicos[dia_nome] = []
                    
                    horario_str = f"{row['horario_inicio']} - {row['horario_fim']}" if row['horario_inicio'] and row['horario_fim'] else "A definir"
                    
                    horarios_aulas_dinamicos[dia_nome].append({
                        'horario': horario_str,
                        'disciplina': row['disciplina'],
                        'professor': row['professor'] or '',
                        'local': row['local'] or ''
                    })
                
                print(f"[DASHBOARD] Usando {len(horarios_rows)} horários do scraper")
            else:
                # Fallback para dados fixos se não houver dados scraped
                horarios_aulas_dinamicos = HORARIOS_AULAS
                print("[DASHBOARD] Usando horários fixos (fallback)")
                
        except Exception as e:
            print(f"[DASHBOARD] Erro ao buscar horários: {e}")
            horarios_aulas_dinamicos = HORARIOS_AULAS

    return render_template(
        'dashboard.html',
        nome=nome,
        curso=curso,
        dark_mode=dark_mode,
        posts_ia=posts_por_curso.get('ia', []),
        posts_ads=posts_por_curso.get('ads', []),
        posts_es=posts_por_curso.get('es', []),
        horarios_aulas=horarios_aulas_dinamicos,
        eventos_academicos=json.dumps(EVENTOS_ACADEMICOS),
        notas=notas,
        faltas=faltas
    )


 


# ============================================
# ROTA DO CHAT
# ============================================
@app.route('/chat', methods=['POST'])
def chat():
    if 'user_id' not in session:
        return jsonify({'error': 'Não autorizado'}), 401
    user_id = session['user_id']

    if not can_make_request(user_id):
        wait_time = get_wait_time(user_id)
        return jsonify({
            'response': f'⏳ Aguarde {wait_time}s.',
            'wait_time': wait_time,
            'rate_limited': True
        })

    data = request.get_json(silent=True) or {}
    mensagem_usuario = sanitizar_html(data.get('message', '').strip())

    resposta_evento = criar_evento_rapido_via_chat(user_id, mensagem_usuario)
    if resposta_evento is not None:
        salvar_historico_chat(user_id, mensagem_usuario, resposta_evento)
        return jsonify({'response': resposta_evento})

    with get_db_connection() as conn:
        c = conn.cursor()
        c.execute('SELECT nome, matricula, email, curso FROM usuarios WHERE id = ?', (user_id,))
        usuario_dados = c.fetchone()

        c.execute(''' 
            SELECT mensagem, resposta  
            FROM historico_chat  
            WHERE usuario_id = ?  
            ORDER BY data_hora DESC  
            LIMIT 6 
        ''', (user_id,))
        historico_recente = c.fetchall()

        c.execute('SELECT disciplina, conteudo_texto FROM conteudos_ava WHERE usuario_id=?', (user_id,))
        conteudos_ava = c.fetchall()

    termo_lower = mensagem_usuario.lower()
    ava_texto = ""

    termos_cal = ['quando', 'data', 'dia', 'feriado', 'provas', 'calendário', 'calendario', 'va1', 'va2']
    quer_calendario = any(t in termo_lower for t in termos_cal)

    contexto_cal = ""
    if quer_calendario:
        contexto_cal = f"CALENDÁRIO:\n{DADOS_ACADEMICOS}\n"

    termos_horarios = ['aula', 'horário', 'horario', 'segunda', 'terça', 'terca',
                       'quarta', 'quinta', 'sexta', 'grade', 'disciplina',
                       'professor', 'hoje', 'amanhã', 'amanha']
    quer_horarios = any(t in termo_lower for t in termos_horarios)

    contexto_horarios = ""
    if quer_horarios:
        contexto_horarios = "\n--- GRADE DE HORÁRIOS ---\n"
        for dia, aulas in HORARIOS_AULAS.items():
            contexto_horarios += f"\n{dia}:\n"
            for aula in aulas:
                contexto_horarios += f"  • {aula['horario']} - {aula['disciplina']}\n"
                contexto_horarios += f"    Professor: {aula['professor']}\n"
        contexto_horarios += "\n"

    match_semana = re.search(r'semana\s*(\d+)', termo_lower)
    semana_foco = match_semana.group(1) if match_semana else None

    materiais_encontrados = False

    if conteudos_ava:
        ava_texto = "--- CONTEÚDOS DETALHADOS DO AVA ---\n"

        for item in conteudos_ava:
            disc_texto = item['conteudo_texto']
            disc_nome = item['disciplina']

            if semana_foco:
                padroes = [f"SEMANA {semana_foco}", f"Semana {semana_foco}", f"Fase {semana_foco}"]

                for padrao in padroes:
                    if padrao in disc_texto:
                        partes = disc_texto.split(padrao)
                        if len(partes) > 1:
                            conteudo_semana = partes[1][:5000]
                            ava_texto += f"\n>>> [{disc_nome}] - {padrao} <<<\n{conteudo_semana}\n"
                            materiais_encontrados = True
                            break

            elif any(t in disc_nome.lower() for t in termo_lower.split() if len(t) > 3):
                ava_texto += f"\n>>> {disc_nome} (Visão Geral) <<<\n{disc_texto[:3000]}\n"
                materiais_encontrados = True

        if not materiais_encontrados:
            if semana_foco:
                ava_texto += f"\nAVISO: Não encontrei detalhes específicos para a Semana {semana_foco} nos textos baixados.\n"

            ava_texto += "Resumo dos materiais disponíveis no banco:\n"
            for item in conteudos_ava:
                ava_texto += f"## {item['disciplina']}:\n{item['conteudo_texto'][:500]}...\n"
    else:
        ava_texto = "AVA não sincronizado."

    historico_texto = ""
    if historico_recente:
        for h in reversed(historico_recente):
            historico_texto += f"Aluno: {h['mensagem']}\nAssistente: {h['resposta']}\n"
    else:
        historico_texto = "Nenhuma conversa anterior."

    if usuario_dados:
        nome_usuario = usuario_dados['nome']
        matricula_usuario = usuario_dados['matricula']
        email_usuario = usuario_dados['email']
        curso_usuario = usuario_dados['curso']
    else:
        nome_usuario = session.get('user_nome', 'Usuário')
        matricula_usuario = 'N/A'
        email_usuario = 'N/A'
        curso_usuario = 'N/A'

    try:
        prompt = f"""
Você é o **IAUniev Professor**, assistente acadêmico da UniEvangélica.

ALUNO: {nome_usuario} ({curso_usuario})

{contexto_cal}
{contexto_horarios}

CONTEÚDO DO AVA (EXTRAÍDO):
{ava_texto}

HISTÓRICO:
{historico_texto}

PERGUNTA:
\"\"\"{mensagem_usuario}\"\"\"

### DIRETRIZES DE RESPOSTA ###
1. **LEIA O TEXTO DO AVA:** O texto acima contém descrições das aulas, resumos e links. Use isso como base.
2. **SEMANA X:** Se o aluno perguntou sobre "Semana X", RESUMA o que está escrito nos tópicos dessa semana. Diga "O conteúdo aborda..." ou "Os tópicos são...". Não diga apenas que "tem atividades".
3. **VÍDEOS E ARQUIVOS:** Se houver links no texto (ex: 🎬 [VIDEOAULA] ou 📎 [PDF]), LISTE-OS EXPLICITAMENTE para o aluno clicar.
   - Exemplo: "📺 Assista à videoaula: [Link]"
   - Exemplo: "📄 Leia o PDF: [Link]"
4. Se o texto extraído não tiver descrição (estiver vazio ou só com títulos genéricos), seja honesto: "O professor não colocou descrição detalhada no AVA, apenas os títulos das atividades."
5. Seja útil e incentive o estudo.
"""

        safety_settings = [
            {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"}
        ]

        if model:
            response = model.generate_content(
                prompt,
                safety_settings=safety_settings
            )
            if response.parts:
                resposta_final = response.text
                salvar_historico_chat(user_id, mensagem_usuario, resposta_final)
                return jsonify({'response': resposta_final})
            else:
                resposta = "Não consegui gerar resposta."
                salvar_historico_chat(user_id, mensagem_usuario, resposta)
                return jsonify({'response': resposta})
        else:
            return jsonify({'response': "Erro de configuração."})

    except Exception as e:
        print(f"Erro Gemini: {str(e)}")
        resposta = "Erro de conexão."
        salvar_historico_chat(user_id, mensagem_usuario, resposta)
        return jsonify({'response': resposta})


# ============================================
# 🆕 V5.1: NOVAS ROTAS PARA SINCRONIZAÇÃO MANUAL AVA
# ============================================

@app.route('/api/status_sync', methods=['GET'])
def status_sync():
    """V5.1: Retorna status da sincronização AVA do usuário"""
    try:
        user_id = session.get('user_id')

        if not user_id:
            return jsonify({'erro': 'Não autenticado'}), 401

        if not SCRAPER_DISPONIVEL:
            return jsonify({
                'erro': 'Scraper não disponível',
                'sincronizando': False,
                'tem_dados': False
            }), 500

        sincronizando = sincronizacoes_em_andamento.get(user_id, False)
        tem_dados = usuario_tem_cache(user_id)
        ultima_sync = obter_ultima_sincronizacao(user_id)

        response = {
            'sincronizando': sincronizando,
            'tem_dados': tem_dados,
            'ultima_sync': None,
            'ultima_sync_formatada': None
        }

        if ultima_sync:
            response['ultima_sync'] = ultima_sync.isoformat()
            response['ultima_sync_formatada'] = ultima_sync.strftime('%d/%m/%Y às %H:%M')

        return jsonify(response), 200

    except Exception as e:
        print(f"[API] Erro ao obter status: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'erro': str(e)}), 500


@app.route('/api/sincronizar_ava', methods=['POST'])
def api_sincronizar_ava():
    """V5.1: Inicia sincronização MANUAL com o AVA"""
    try:
        user_id = session.get('user_id')

        if not user_id:
            return jsonify({'erro': 'Não autenticado'}), 401

        if not SCRAPER_DISPONIVEL:
            return jsonify({'erro': 'Scraper não disponível'}), 500

        if sincronizacoes_em_andamento.get(user_id, False):
            print(f"[API] Sincronização já em andamento para user {user_id}")
            return jsonify({
                'status': 'em_andamento',
                'mensagem': 'Sincronização já está em andamento'
            }), 200

        matricula = session.get('matricula')
        cpf = session.get('cpf')
        nome = session.get('user_nome')

        if not matricula or not cpf:
            print(f"[API] Erro: Credenciais não encontradas na sessão para user {user_id}")
            return jsonify({
                'erro': 'Dados de login não encontrados na sessão'
            }), 400

        print(f"[API] Iniciando sincronização para: {nome} (ID: {user_id}, Matrícula: {matricula})")

        thread = threading.Thread(
            target=executar_sincronizacao_monitorada,
            args=(user_id, matricula, cpf, True),
            daemon=True
        )
        thread.start()

        return jsonify({
            'status': 'iniciado',
            'mensagem': 'Sincronização iniciada com sucesso',
            'estimativa': '5-8 minutos',
            'user_id': user_id
        }), 200

    except Exception as e:
        print(f"[API] Erro ao iniciar sincronização: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'erro': str(e)}), 500


@app.route('/api/cancelar_sync', methods=['POST'])
def cancelar_sync():
    """V5.1: Marca sincronização como cancelada"""
    try:
        user_id = session.get('user_id')

        if not user_id:
            return jsonify({'erro': 'Não autenticado'}), 401

        if user_id in sincronizacoes_em_andamento:
            sincronizacoes_em_andamento[user_id] = False
            status_sincronizacao[user_id] = 'cancelado'
            print(f"[API] Sincronização marcada como cancelada para user {user_id}")

        return jsonify({
            'status': 'cancelado',
            'mensagem': 'Status atualizado'
        }), 200

    except Exception as e:
        print(f"[API] Erro ao cancelar: {e}")
        return jsonify({'erro': str(e)}), 500


# ============================================
# 🆕 ROTAS API - LYCEUM V5.1
# ============================================

@app.route('/api/status_sync_lyceum', methods=['GET'])
def status_sync_lyceum():
    """V5.1: Retorna status da sincronização Lyceum"""
    try:
        user_id = session.get('user_id')

        if not user_id:
            return jsonify({'erro': 'Não autenticado'}), 401

        if not LYCEUM_DISPONIVEL:
            return jsonify({
                'erro': 'Scraper Lyceum não disponível',
                'sincronizando': False,
                'tem_dados': False
            }), 500

        sincronizando = sincronizacoes_em_andamento_lyceum.get(user_id, False)
        tem_dados = usuario_tem_cache_lyceum(user_id)
        ultima_sync = obter_ultima_sincronizacao_lyceum(user_id)

        response = {
            'sincronizando': sincronizando,
            'tem_dados': tem_dados,
            'ultima_sync': None,
            'ultima_sync_formatada': None
        }

        if ultima_sync:
            response['ultima_sync'] = ultima_sync.isoformat()
            response['ultima_sync_formatada'] = ultima_sync.strftime('%d/%m/%Y às %H:%M')

        return jsonify(response), 200

    except Exception as e:
        print(f"[API LYCEUM] Erro ao obter status: {e}")
        return jsonify({'erro': str(e)}), 500


@app.route('/api/sincronizar_lyceum', methods=['POST'])
def api_sincronizar_lyceum():
    """V5.1: Inicia sincronização MANUAL com Lyceum"""
    try:
        user_id = session.get('user_id')

        if not user_id:
            return jsonify({'erro': 'Não autenticado'}), 401

        if not LYCEUM_DISPONIVEL:
            return jsonify({'erro': 'Scraper Lyceum não disponível'}), 500

        if sincronizacoes_em_andamento_lyceum.get(user_id, False):
            print(f"[API LYCEUM] Sincronização já em andamento para user {user_id}")
            return jsonify({
                'status': 'em_andamento',
                'mensagem': 'Sincronização já está em andamento'
            }), 200

        matricula = session.get('matricula')
        nome = session.get('user_nome')

        # Tenta obter senha_lyceum do banco
        senha_lyceum = None
        try:
            with get_db_connection() as conn:
                c = conn.cursor()
                c.execute('SELECT senha_lyceum, cpf FROM usuarios WHERE id = ?', (user_id,))
                row = c.fetchone()
                if row:
                    senha_lyceum = row['senha_lyceum'] or row['cpf']
        except:
            pass

        if not matricula or not senha_lyceum:
            print(f"[API LYCEUM] Erro: Credenciais não encontradas para user {user_id}")
            return jsonify({
                'erro': 'Credenciais não encontradas. Configure sua senha do Lyceum.'
            }), 400

        print(f"[API LYCEUM] Iniciando sincronização para: {nome} (ID: {user_id})")

        thread = threading.Thread(
            target=executar_sincronizacao_monitorada_lyceum,
            args=(user_id, matricula, senha_lyceum, True),
            daemon=True
        )
        thread.start()

        return jsonify({
            'status': 'iniciado',
            'mensagem': 'Sincronização Lyceum iniciada com sucesso',
            'estimativa': '2-3 minutos',
            'user_id': user_id
        }), 200

    except Exception as e:
        print(f"[API LYCEUM] Erro ao iniciar sincronização: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'erro': str(e)}), 500


@app.route('/api/cancelar_sync_lyceum', methods=['POST'])
def cancelar_sync_lyceum():
    """V5.1: Marca sincronização Lyceum como cancelada"""
    try:
        user_id = session.get('user_id')

        if not user_id:
            return jsonify({'erro': 'Não autenticado'}), 401

        if user_id in sincronizacoes_em_andamento_lyceum:
            sincronizacoes_em_andamento_lyceum[user_id] = False
            status_sincronizacao_lyceum[user_id] = 'cancelado'
            print(f"[API LYCEUM] Sincronização marcada como cancelada para user {user_id}")

        return jsonify({
            'status': 'cancelado',
            'mensagem': 'Status atualizado'
        }), 200

    except Exception as e:
        print(f"[API LYCEUM] Erro ao cancelar: {e}")
        return jsonify({'erro': str(e)}), 500


# ============================================
# ROTAS API - COMUNIDADE
# ============================================
@app.route('/api/historico_chat')
def buscar_historico_chat():
    if 'user_id' not in session:
        return jsonify({'error': 'Não autorizado'}), 401
    user_id = session['user_id']
    limite = request.args.get('limite', 20, type=int)
    try:
        with get_db_connection() as conn:
            c = conn.cursor()
            c.execute(''' 
                SELECT mensagem, resposta, strftime('%d/%m/%Y %H:%M', data_hora) as data_formatada 
                FROM historico_chat 
                WHERE usuario_id = ? 
                ORDER BY data_hora DESC 
                LIMIT ? 
            ''', (user_id, limite))
            historico = []
            for row in c.fetchall():
                historico.append({
                    'mensagem': row['mensagem'],
                    'resposta': row['resposta'],
                    'data': row['data_formatada']
                })
        return jsonify({'success': True, 'historico': historico})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/limpar_historico', methods=['POST'])
def limpar_historico():
    if 'user_id' not in session:
        return jsonify({'error': 'Não autorizado'}), 401
    user_id = session['user_id']
    try:
        with get_db_connection() as conn:
            conn.execute('DELETE FROM historico_chat WHERE usuario_id = ?', (user_id,))
            conn.commit()
        return jsonify({'success': True, 'message': 'Histórico limpo!'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/posts')
def buscar_posts():
    """Lazy loading de posts"""
    if 'user_id' not in session:
        return jsonify({'error': 'Não autorizado'}), 401

    curso = request.args.get('curso', 'ia')
    page = request.args.get('page', 1, type=int)
    limit = request.args.get('limit', 10, type=int)
    offset = (page - 1) * limit

    user_id = session['user_id']

    try:
        with get_db_connection() as conn:
            c = conn.cursor()
            c.execute('''
                SELECT post_id, tipo, titulo, conteudo, tags, nome_usuario,
                       strftime('%d/%m/%Y %H:%M', data_criacao) as data_formatada,
                       usuario_id
                FROM posts
                WHERE curso = ?
                ORDER BY data_criacao DESC
                LIMIT ? OFFSET ?
            ''', (curso, limit, offset))

            posts = []
            for row in c.fetchall():
                posts.append({
                    'post_id': row['post_id'],
                    'tipo': row['tipo'],
                    'titulo': row['titulo'],
                    'conteudo': row['conteudo'],
                    'tags': row['tags'] or '',
                    'nome_usuario': row['nome_usuario'],
                    'data': row['data_formatada'],
                    'e_meu': row['usuario_id'] == user_id
                })

        return jsonify({'success': True, 'posts': posts})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/criar_post', methods=['POST'])
def criar_post():
    if 'user_id' not in session:
        return jsonify({'error': 'Não autorizado'}), 401
    data = request.get_json(silent=True) or {}
    curso = data.get('curso')
    tipo = data.get('tipo')
    titulo = sanitizar_html(data.get('titulo'))
    conteudo = sanitizar_html(data.get('conteudo'))
    tags = sanitizar_html(data.get('tags', ''))
    user_id = session['user_id']
    nome_usuario = session['user_nome']
    post_id = f"post-{curso}-{int(datetime.now().timestamp())}"

    try:
        with get_db_connection() as conn:
            conn.execute(''' 
                INSERT INTO posts (post_id, curso, tipo, titulo, conteudo, tags, usuario_id, nome_usuario) 
                VALUES (?, ?, ?, ?, ?, ?, ?, ?) 
            ''', (post_id, curso, tipo, titulo, conteudo, tags, user_id, nome_usuario))
            conn.commit()
        return jsonify({
            'success': True,
            'post_id': post_id,
            'nome_usuario': nome_usuario,
            'data': datetime.now().strftime('%d/%m/%Y %H:%M')
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/curtir', methods=['POST'])
def curtir_post():
    if 'user_id' not in session:
        return jsonify({'error': 'Não autorizado'}), 401
    data = request.get_json(silent=True) or {}
    post_id = data.get('post_id')
    user_id = session['user_id']

    try:
        with get_db_connection() as conn:
            c = conn.cursor()
            c.execute('SELECT * FROM curtidas WHERE post_id = ? AND usuario_id = ?', (post_id, user_id))
            curtida_existe = c.fetchone()

            c.execute('SELECT usuario_id, titulo FROM posts WHERE post_id = ?', (post_id,))
            post = c.fetchone()

            if curtida_existe:
                c.execute('DELETE FROM curtidas WHERE post_id = ? AND usuario_id = ?', (post_id, user_id))
                acao = 'descurtiu'
            else:
                c.execute('INSERT INTO curtidas (post_id, usuario_id) VALUES (?, ?)', (post_id, user_id))
                acao = 'curtiu'

                if post and post['usuario_id'] != user_id:
                    criar_notificacao(
                        post['usuario_id'],
                        'curtida',
                        f"{session['user_nome']} curtiu seu post: {post['titulo']}",
                        f"/dashboard#post-{post_id}"
                    )

            c.execute('SELECT COUNT(*) AS total FROM curtidas WHERE post_id = ?', (post_id,))
            total_curtidas = c.fetchone()['total']
            conn.commit()
        return jsonify({'success': True, 'acao': acao, 'total_curtidas': total_curtidas})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/comentar', methods=['POST'])
def comentar_post():
    if 'user_id' not in session:
        return jsonify({'error': 'Não autorizado'}), 401
    data = request.get_json(silent=True) or {}
    post_id = data.get('post_id')
    comentario = sanitizar_html(data.get('comentario'))
    user_id = session['user_id']
    nome_usuario = session['user_nome']

    try:
        with get_db_connection() as conn:
            c = conn.cursor()

            c.execute('SELECT usuario_id, titulo FROM posts WHERE post_id = ?', (post_id,))
            post = c.fetchone()

            c.execute(''' 
                INSERT INTO comentarios (post_id, usuario_id, nome_usuario, comentario) 
                VALUES (?, ?, ?, ?) 
            ''', (post_id, user_id, nome_usuario, comentario))
            comentario_id = c.lastrowid
            conn.commit()

            if post and post['usuario_id'] != user_id:
                criar_notificacao(
                    post['usuario_id'],
                    'comentario',
                    f"{nome_usuario} comentou no seu post: {post['titulo']}",
                    f"/dashboard#post-{post_id}"
                )

        return jsonify({
            'success': True,
            'comentario_id': comentario_id,
            'nome_usuario': nome_usuario,
            'comentario': comentario,
            'data': datetime.now().strftime('%d/%m/%Y %H:%M')
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/excluir_post', methods=['POST'])
def excluir_post():
    if 'user_id' not in session:
        return jsonify({'error': 'Não autorizado'}), 401
    data = request.get_json(silent=True) or {}
    post_id = data.get('post_id')
    user_id = session['user_id']

    try:
        with get_db_connection() as conn:
            c = conn.cursor()
            c.execute('SELECT usuario_id FROM posts WHERE post_id = ?', (post_id,))
            post = c.fetchone()
            if not post:
                return jsonify({'error': 'Post não encontrado'}), 404
            if post['usuario_id'] != user_id:
                return jsonify({'error': 'Sem permissão'}), 403
            c.execute('DELETE FROM comentarios WHERE post_id = ?', (post_id,))
            c.execute('DELETE FROM curtidas WHERE post_id = ?', (post_id,))
            c.execute('DELETE FROM posts WHERE post_id = ?', (post_id,))
            conn.commit()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/excluir_comentario', methods=['POST'])
def excluir_comentario():
    if 'user_id' not in session:
        return jsonify({'error': 'Não autorizado'}), 401
    data = request.get_json(silent=True) or {}
    comentario_id = data.get('comentario_id')
    user_id = session['user_id']

    try:
        with get_db_connection() as conn:
            c = conn.cursor()
            c.execute('SELECT usuario_id FROM comentarios WHERE id = ?', (comentario_id,))
            comentario = c.fetchone()
            if not comentario:
                return jsonify({'error': 'Comentário não encontrado'}), 404
            if comentario['usuario_id'] != user_id:
                return jsonify({'error': 'Sem permissão'}), 403
            c.execute('DELETE FROM comentarios WHERE id = ?', (comentario_id,))
            conn.commit()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/interacoes/<post_id>')
def buscar_interacoes(post_id):
    if 'user_id' not in session:
        return jsonify({'error': 'Não autorizado'}), 401
    user_id = session['user_id']

    try:
        with get_db_connection() as conn:
            c = conn.cursor()
            c.execute('SELECT COUNT(*) AS total FROM curtidas WHERE post_id = ?', (post_id,))
            total_curtidas = c.fetchone()['total']
            c.execute('SELECT * FROM curtidas WHERE post_id = ? AND usuario_id = ?', (post_id, user_id))
            usuario_curtiu = c.fetchone() is not None
            c.execute(''' 
                SELECT id, nome_usuario, comentario, strftime('%d/%m/%Y %H:%M', data_comentario) as data_formatada, usuario_id 
                FROM comentarios 
                WHERE post_id = ? 
                ORDER BY data_comentario ASC 
            ''', (post_id,))
            comentarios = []
            for row in c.fetchall():
                comentarios.append({
                    'id': row['id'],
                    'nome_usuario': row['nome_usuario'],
                    'comentario': row['comentario'],
                    'data': row['data_formatada'],
                    'e_meu': row['usuario_id'] == user_id
                })
        return jsonify({
            'success': True,
            'total_curtidas': total_curtidas,
            'usuario_curtiu': usuario_curtiu,
            'total_comentarios': len(comentarios),
            'comentarios': comentarios
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ============================================
# ROTAS DO CALENDÁRIO
# ============================================
@app.route('/api/eventos_calendario')
def buscar_eventos_calendario():
    if 'user_id' not in session:
        return jsonify({'error': 'Não autorizado'}), 401
    user_id = session['user_id']

    try:
        timestamp = int(time.time() // 300)
        eventos_db = get_eventos_cache_key(user_id, timestamp)

        eventos_pessoais = []
        for row in eventos_db:
            start_datetime = row['data_evento']
            if row['hora_evento']:
                start_datetime += 'T' + row['hora_evento']

            eventos_pessoais.append({
                'id': row['id'],
                'title': row['titulo'],
                'start': start_datetime,
                'color': row['cor'],
                'extendedProps': {
                    'descricao': row['descricao'],
                    'tipo': row['tipo'],
                    'pessoal': True,
                    'alerta': row['alerta'],
                    'minutos_antes_alerta': row['minutos_antes_alerta']
                },
                'allDay': False if row['hora_evento'] else True
            })

        # ============================================
        # 🆕 V7.0: BUSCAR EVENTOS DO CALENDÁRIO LYCEUM
        # ============================================
        eventos_lyceum = []
        try:
            with get_db_connection() as conn:
                c = conn.cursor()
                c.execute('''
                    SELECT id, titulo, data_evento, tipo, cor, descricao
                    FROM calendario_lyceum
                    WHERE usuario_id = ?
                ''', (user_id,))
                
                for row in c.fetchall():
                    eventos_lyceum.append({
                        'id': f"lyceum_{row['id']}",
                        'title': row['titulo'],
                        'start': row['data_evento'],
                        'color': row['cor'] or '#4a90e2',
                        'extendedProps': {
                            'descricao': row['descricao'] or '',
                            'tipo': row['tipo'] or 'evento',
                            'pessoal': False,
                            'origem': 'lyceum'
                        },
                        'allDay': True
                    })
                
                if eventos_lyceum:
                    print(f"[CALENDARIO] Carregados {len(eventos_lyceum)} eventos do Lyceum")
        except Exception as e:
            # Tabela pode não existir ainda
            print(f"[CALENDARIO] Eventos Lyceum não disponíveis: {e}")

        # Combinar: Eventos fixos + Lyceum + Pessoais
        # Se tiver eventos do Lyceum, usar eles como base. Senão, usar fixos.
        if eventos_lyceum:
            todos_eventos = eventos_lyceum + eventos_pessoais
        else:
            todos_eventos = EVENTOS_ACADEMICOS + eventos_pessoais
            
        return jsonify({'success': True, 'eventos': todos_eventos})
    except Exception as e:
        print(f"[ERROR] Erro: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/criar_evento', methods=['POST'])
def criar_evento():
    if 'user_id' not in session:
        return jsonify({'error': 'Não autorizado'}), 401
    user_id = session['user_id']
    data = request.get_json(silent=True) or {}

    titulo = sanitizar_html(data.get('titulo'))
    descricao = sanitizar_html(data.get('descricao', ''))
    data_evento = data.get('data')
    hora_evento = data.get('hora', '')
    tipo = data.get('tipo', 'pessoal')
    cor = data.get('cor', '#4a90e2')
    alerta = data.get('alerta', 0)
    minutos_antes_alerta = data.get('minutos_antes_alerta', 30)

    if not titulo or not data_evento:
        return jsonify({'error': 'Título e data obrigatórios!'}), 400

    try:
        with get_db_connection() as conn:
            c = conn.cursor()
            c.execute(''' 
                INSERT INTO eventos_calendario (usuario_id, titulo, descricao, data_evento, hora_evento, tipo, cor, alerta, minutos_antes_alerta) 
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?) 
            ''', (user_id, titulo, descricao, data_evento, hora_evento, tipo, cor, alerta, minutos_antes_alerta))
            evento_id = c.lastrowid
            conn.commit()

        limpar_cache_eventos()

        return jsonify({
            'success': True,
            'evento_id': evento_id,
            'mensagem': 'Evento criado!'
        })
    except Exception as e:
        print(f"[ERROR] Erro: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/excluir_evento', methods=['POST'])
def excluir_evento():
    if 'user_id' not in session:
        return jsonify({'error': 'Não autorizado'}), 401
    user_id = session['user_id']
    data: dict = request.get_json(silent=True) or {}
    evento_id = data.get('evento_id')

    try:
        with get_db_connection() as conn:
            c = conn.cursor()
            c.execute('SELECT usuario_id FROM eventos_calendario WHERE id = ?', (evento_id,))
            evento = c.fetchone()
            if not evento:
                return jsonify({'error': 'Evento não encontrado'}), 404
            if evento['usuario_id'] != user_id:
                return jsonify({'error': 'Sem permissão'}), 403
            c.execute('DELETE FROM eventos_calendario WHERE id = ?', (evento_id,))
            conn.commit()

        limpar_cache_eventos()

        return jsonify({'success': True, 'mensagem': 'Evento excluído!'})
    except Exception as e:
        print(f"[ERROR] Erro: {e}")
        return jsonify({'error': str(e)}), 500


# ============================================
# ROTAS DE STATUS
# ============================================
@app.route('/api/checar_status_sync')
def checar_status_sync():
    """Mantida para compatibilidade com tela de loading"""
    if 'user_id' not in session:
        return jsonify({'status': 'erro', 'redirect': url_for('index')})

    user_id = session['user_id']
    estado = status_sincronizacao.get(user_id, 'concluido')

    return jsonify({'status': estado})


@app.route('/api/status_sincronizacao')
def status_sincronizacao_dashboard():
    """Mantida para compatibilidade"""
    if 'user_id' not in session:
        return jsonify({'error': 'Não autorizado'}), 401

    user_id = session['user_id']
    with get_db_connection() as conn:
        c = conn.cursor()
        c.execute('SELECT COUNT(*) FROM conteudos_ava WHERE usuario_id = ?', (user_id,))
        count_conteudos = c.fetchone()[0]

    sincronizado = count_conteudos > 0
    detalhes = f"{count_conteudos} materiais baixados."

    return jsonify({
        'sincronizado': sincronizado,
        'detalhes': detalhes
    })


# ============================================
# MODO ESCURO
# ============================================
@app.route('/api/toggle_dark_mode', methods=['POST'])
def toggle_dark_mode():
    if 'user_id' not in session:
        return jsonify({'error': 'Não autorizado'}), 401

    user_id = session['user_id']

    try:
        with get_db_connection() as conn:
            c = conn.cursor()
            c.execute('SELECT dark_mode FROM usuarios WHERE id = ?', (user_id,))
            current = c.fetchone()['dark_mode']
            new_mode = 0 if current == 1 else 1

            c.execute('UPDATE usuarios SET dark_mode = ? WHERE id = ?', (new_mode, user_id))
            conn.commit()

            session['dark_mode'] = new_mode

        return jsonify({'success': True, 'dark_mode': new_mode})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ============================================
# EXPORT DE DADOS
# ============================================
@app.route('/api/exportar_notas/<formato>')
def exportar_notas(formato):
    if 'user_id' not in session:
        return jsonify({'error': 'Não autorizado'}), 401

    user_id = session['user_id']

    with get_db_connection() as conn:
        c = conn.cursor()
        c.execute('''SELECT disciplina, va1, va2, va3, media, situacao 
                     FROM notas_aluno WHERE usuario_id = ?''', (user_id,))
        notas = c.fetchall()

        c.execute('SELECT nome FROM usuarios WHERE id = ?', (user_id,))
        nome = c.fetchone()['nome']

    if formato == 'pdf':
        buffer = io.BytesIO()
        p: Any = Canvas(buffer, pagesize=A4)

        y = 800
        p.setFont("Helvetica-Bold", 16)
        p.drawString(100, y, "RELATÓRIO DE NOTAS - IAUniev")

        y -= 30
        p.setFont("Helvetica", 12)
        p.drawString(100, y, f"Aluno: {nome}")

        y -= 30
        p.drawString(100, y, f"Data: {datetime.now().strftime('%d/%m/%Y')}")

        y -= 50

        data_table = [['Disciplina', 'VA1', 'VA2', 'VA3', 'Média', 'Situação']]
        for nota in notas:
            data_table.append([
                nota['disciplina'][:30],
                str(nota['va1']),
                str(nota['va2']),
                str(nota['va3']),
                str(nota['media']),
                nota['situacao']
            ])

        table = Table(data_table, colWidths=[200, 50, 50, 50, 50, 80])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))

        table.wrapOn(p, 100, y)
        table.drawOn(p, 50, y - len(data_table) * 25)

        p.save()
        buffer.seek(0)

        return send_file(buffer,
                         mimetype='application/pdf',
                         as_attachment=True,
                         download_name=f'notas_{nome}.pdf')

    elif formato == 'excel':
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment

            wb = openpyxl.Workbook()
            ws = wb.active
            if ws is None:
                ws = wb.create_sheet("Notas")
            ws.title = "Notas"

            ws['A1'] = f"Relatório de Notas - {nome}"
            ws['A1'].font = Font(size=14, bold=True)
            ws['A2'] = f"Data: {datetime.now().strftime('%d/%m/%Y')}"

            headers = ['Disciplina', 'VA1', 'VA2', 'VA3', 'Média', 'Situação']
            for col, header in enumerate(headers, start=1):
                cell = ws.cell(row=4, column=col, value=header)
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
                cell.alignment = Alignment(horizontal="center")

            for row_idx, nota in enumerate(notas, start=5):
                ws.cell(row=row_idx, column=1, value=nota['disciplina'])
                ws.cell(row=row_idx, column=2, value=nota['va1'])
                ws.cell(row=row_idx, column=3, value=nota['va2'])
                ws.cell(row=row_idx, column=4, value=nota['va3'])
                ws.cell(row=row_idx, column=5, value=nota['media'])
                ws.cell(row=row_idx, column=6, value=nota['situacao'])

            ws.column_dimensions['A'].width = 40
            for col in ['B', 'C', 'D', 'E']:
                ws.column_dimensions[col].width = 10
            ws.column_dimensions['F'].width = 15

            buffer = io.BytesIO()
            wb.save(buffer)
            buffer.seek(0)

            return send_file(buffer,
                             mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                             as_attachment=True,
                             download_name=f'notas_{nome}.xlsx')
        except ImportError:
            return jsonify({'error': 'openpyxl não instalado'}), 500

    return jsonify({'error': 'Formato inválido'}), 400


from flask import request, jsonify
from werkzeug.utils import secure_filename
import os
import PyPDF2
import docx
from youtube_transcript_api import YouTubeTranscriptApi
import google.generativeai as genai
from dotenv import load_dotenv

# Carregar variáveis de ambiente do .env
load_dotenv()

# Configurar Gemini com a chave do .env
genai.configure(api_key=os.getenv('GEMINI_API_KEY'))

# Configurações de upload
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'txt', 'doc', 'docx'}

# Criar pasta de uploads se não existir
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)


def allowed_file(filename):
    """Verifica se o arquivo tem extensão permitida"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_pdf(filepath):
    """Extrai texto de um PDF - SEM LIMITES"""
    try:
        text = ""
        with open(filepath, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            total_pages = len(pdf_reader.pages)
            print(f"[PDF] Processando {total_pages} páginas...")

            for i, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    text += page_text + "\n"
                    if (i + 1) % 10 == 0:
                        print(f"[PDF] Processadas {i + 1}/{total_pages} páginas")
                except Exception as e:
                    print(f"[PDF] Erro na página {i + 1}: {e}")
                    continue

        print(f"[PDF] ✅ Extraído {len(text)} caracteres")
        return text.strip()
    except Exception as e:
        print(f"[PDF] ❌ Erro ao extrair texto: {e}")
        return None


def extract_text_from_docx(filepath):
    """Extrai texto de um DOCX - SEM LIMITES"""
    try:
        doc = docx.Document(filepath)
        text = ""
        total_paragraphs = len(doc.paragraphs)
        print(f"[DOCX] Processando {total_paragraphs} parágrafos...")

        for i, paragraph in enumerate(doc.paragraphs):
            text += paragraph.text + "\n"
            if (i + 1) % 100 == 0:
                print(f"[DOCX] Processados {i + 1}/{total_paragraphs} parágrafos")

        print(f"[DOCX] ✅ Extraído {len(text)} caracteres")
        return text.strip()
    except Exception as e:
        print(f"[DOCX] ❌ Erro ao extrair texto: {e}")
        return None


def extract_text_from_txt(filepath):
    """Extrai texto de um TXT - SEM LIMITES"""
    try:
        with open(filepath, 'r', encoding='utf-8') as file:
            text = file.read()
        print(f"[TXT] ✅ Extraído {len(text)} caracteres")
        return text.strip()
    except UnicodeDecodeError:
        # Tenta com outra codificação
        try:
            with open(filepath, 'r', encoding='latin-1') as file:
                text = file.read()
            print(f"[TXT] ✅ Extraído {len(text)} caracteres (latin-1)")
            return text.strip()
        except Exception as e:
            print(f"[TXT] ❌ Erro ao ler arquivo: {e}")
            return None
    except Exception as e:
        print(f"[TXT] ❌ Erro ao ler arquivo: {e}")
        return None


# ============================================
# FUNÇÃO ALTERNATIVA - YOUTUBE TRANSCRIPT
# Use esta versão se a anterior não funcionar
# ============================================

# ============================================
# SOLUÇÃO DEFINITIVA - YOUTUBE TRANSCRIPT
# Múltiplas abordagens para garantir funcionamento
# ============================================

def extract_youtube_transcript(video_id):
    """
    Extrai transcrição usando yt-dlp
    Método mais robusto e confiável
    """
    import subprocess
    import json
    import os
    import tempfile

    print(f"[YOUTUBE] Video ID: {video_id}")

    try:
        # Verificar se yt-dlp está instalado
        try:
            result = subprocess.run(
                ['yt-dlp', '--version'],
                capture_output=True,
                text=True,
                timeout=5
            )
            if result.returncode != 0:
                raise Exception("yt-dlp não está instalado")
        except FileNotFoundError:
            print("[YOUTUBE] ❌ yt-dlp não encontrado")
            print("[YOUTUBE] Instale: pip install yt-dlp --break-system-packages")
            return None

        # URL do vídeo
        url = f"https://www.youtube.com/watch?v={video_id}"

        # Criar diretório temporário
        with tempfile.TemporaryDirectory() as temp_dir:
            subtitle_file = os.path.join(temp_dir, 'subtitle')

            print("[YOUTUBE] Baixando legendas com yt-dlp...")

            # Comando para baixar legendas
            cmd = [
                'yt-dlp',
                '--write-auto-sub',  # Legendas automáticas
                '--write-sub',  # Legendas normais
                '--sub-lang', 'pt,pt-BR,en',  # Idiomas
                '--skip-download',  # Não baixar vídeo
                '--sub-format', 'vtt',  # Formato
                '-o', subtitle_file,
                url
            ]

            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=30
            )

            # Procurar arquivo de legenda
            subtitle_files = [
                f for f in os.listdir(temp_dir)
                if f.endswith('.vtt')
            ]

            if not subtitle_files:
                print("[YOUTUBE] ❌ Nenhuma legenda baixada")
                print(f"[YOUTUBE] Saída: {result.stderr[:200]}")
                return None

            # Ler primeira legenda encontrada
            subtitle_path = os.path.join(temp_dir, subtitle_files[0])

            print(f"[YOUTUBE] ✅ Legenda encontrada: {subtitle_files[0]}")

            with open(subtitle_path, 'r', encoding='utf-8') as f:
                vtt_content = f.read()

            # Parse VTT
            lines = vtt_content.split('\n')
            texts = []

            for line in lines:
                line = line.strip()
                # Pular linhas vazias, timestamps e metadata
                if not line or '-->' in line or line.startswith('WEBVTT') or line.isdigit():
                    continue
                # Pular tags
                if line.startswith('<') or line.startswith('['):
                    continue

                texts.append(line)

            if not texts:
                print("[YOUTUBE] ❌ Nenhum texto extraído")
                return None

            full_text = " ".join(texts)

            print(f"[YOUTUBE] ✅ {len(full_text)} caracteres, {len(texts)} linhas")
            return full_text

    except subprocess.TimeoutExpired:
        print("[YOUTUBE] ❌ Timeout ao executar yt-dlp")
        return None

    except Exception as e:
        print(f"[YOUTUBE] ❌ Erro: {e}")
        import traceback
        traceback.print_exc()
        return None


# ============================================
# TESTE
# ============================================

if __name__ == "__main__":
    print("🧪 TESTANDO COM YT-DLP")
    print("=" * 70)

    videos = [
        ("dQw4w9WgXcQ", "Never Gonna Give You Up"),
        ("CYRllf5f6HE", "Seu vídeo"),
    ]

    for video_id, title in videos:
        print(f"\n📹 {title}")
        print("-" * 70)

        result = extract_youtube_transcript(video_id)

        if result:
            print(f"\n✅ SUCESSO! {len(result)} caracteres")
            print(f"Preview: {result[:200]}...")
        else:
            print("\n❌ FALHOU")

        print("=" * 70)


def chamar_gemini_api(mensagem):
    """
    Chama a API do Gemini com a mensagem fornecida
    """
    try:
        # Criar modelo Gemini
        model = genai.GenerativeModel('models/gemini-2.5-flash')

        # Gerar resposta
        print(f"[GEMINI] Enviando {len(mensagem)} caracteres para análise...")
        response = model.generate_content(mensagem)

        print(f"[GEMINI] ✅ Resposta recebida")
        return response.text

    except Exception as e:
        print(f"[GEMINI API] ❌ Erro: {e}")
        return f"Erro ao processar sua solicitação: {str(e)}"


# ============================================
# ROTA: Chat com arquivo (SEM LIMITES)
# ============================================
@app.route('/chat/with-file', methods=['POST'])
def chat_with_file():
    """Processa mensagem com arquivo anexado - SEM LIMITES DE TAMANHO"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'Nenhum arquivo enviado'}), 400

        file = request.files['file']
        message = request.form.get('message', '')

        if file.filename == '':
            return jsonify({'error': 'Arquivo sem nome'}), 400

        if not allowed_file(file.filename):
            return jsonify({'error': 'Tipo de arquivo não permitido. Use: PDF, DOC, DOCX ou TXT'}), 400

        print(f"\n[ARQUIVO] Recebido: {file.filename}")

        # Salvar arquivo temporariamente
        filename = secure_filename(cast(str, file.filename))
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        file.save(filepath)

        file_size = os.path.getsize(filepath)
        print(f"[ARQUIVO] Tamanho: {file_size / (1024 * 1024):.2f} MB")

        # Extrair texto baseado no tipo de arquivo
        extension = filename.rsplit('.', 1)[1].lower()

        if extension == 'pdf':
            document_text = extract_text_from_pdf(filepath)
        elif extension == 'docx':
            document_text = extract_text_from_docx(filepath)
        elif extension in ['txt', 'doc']:
            document_text = extract_text_from_txt(filepath)
        else:
            document_text = None

        # Remover arquivo após processamento
        try:
            os.remove(filepath)
            print(f"[ARQUIVO] ✅ Arquivo temporário removido")
        except:
            pass

        if not document_text:
            return jsonify({'error': 'Não foi possível extrair texto do arquivo'}), 400

        # ✅ SEM LIMITE DE TAMANHO - Processa todo o texto
        print(f"[PROCESSAMENTO] Texto extraído: {len(document_text)} caracteres")

        # Criar prompt com contexto do documento
        enhanced_message = f"""Você recebeu um documento para análise.

DOCUMENTO ({len(document_text)} caracteres):
---
{document_text}
---

PERGUNTA DO USUÁRIO: {message if message else "Faça um resumo completo e detalhado deste documento"}

Por favor, responda à pergunta com base no conteúdo completo do documento fornecido acima."""

        print(f"[GEMINI] Enviando {len(enhanced_message)} caracteres para análise...")

        # Chamar Gemini API
        response_text = chamar_gemini_api(enhanced_message)

        print(f"[GEMINI] ✅ Resposta recebida: {len(response_text)} caracteres")

        return jsonify({
            'success': True,
            'response': response_text,
            'document_filename': filename,
            'document_size': len(document_text),
            'rate_limited': False
        })

    except Exception as e:
        print(f"[ERRO] Chat com arquivo: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ============================================
# ROTA: Chat com YouTube (SEM LIMITES)
# ============================================
@app.route('/chat/with-youtube', methods=['POST'])
def chat_with_youtube():
    """Processa mensagem com link do YouTube - SEM LIMITES DE TAMANHO"""
    try:
        data = request.get_json()
        message = data.get('message', '')
        video_id = data.get('video_id', '')

        if not video_id:
            return jsonify({'error': 'ID do vídeo não fornecido'}), 400

        print(f"\n[YOUTUBE] Video ID: {video_id}")

        # Extrair transcrição do YouTube
        transcript_text = extract_youtube_transcript(video_id)

        if not transcript_text:
            return jsonify({
                'error': 'Não foi possível obter a transcrição do vídeo. Possíveis causas:\n' +
                         '• O vídeo não tem legendas\n' +
                         '• As legendas estão desativadas\n' +
                         '• O vídeo é privado ou restrito'
            }), 400

        # ✅ SEM LIMITE DE TAMANHO - Processa toda a transcrição
        print(f"[PROCESSAMENTO] Transcrição: {len(transcript_text)} caracteres")

        # Criar prompt com contexto do vídeo
        enhanced_message = f"""Você recebeu a transcrição completa de um vídeo do YouTube para análise.

TRANSCRIÇÃO DO VÍDEO (ID: {video_id}, {len(transcript_text)} caracteres):
---
{transcript_text}
---

PERGUNTA DO USUÁRIO: {message if message else "Faça um resumo completo e detalhado deste vídeo"}

Por favor, responda à pergunta com base no conteúdo completo da transcrição do vídeo fornecida acima."""

        print(f"[GEMINI] Enviando {len(enhanced_message)} caracteres para análise...")

        # Chamar Gemini API
        response_text = chamar_gemini_api(enhanced_message)

        print(f"[GEMINI] ✅ Resposta recebida: {len(response_text)} caracteres")

        return jsonify({
            'success': True,
            'response': response_text,
            'video_id': video_id,
            'transcript_size': len(transcript_text),
            'rate_limited': False
        })

    except Exception as e:
        print(f"[ERRO] Chat com YouTube: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ============================================
# EXEMPLOS DE USO E COMANDOS
# ============================================
"""
📚 EXEMPLOS DE PERGUNTAS PARA PDFs/DOCUMENTOS:

Resumos:
- "Resuma este documento"
- "Faça um resumo executivo"
- "Quais os pontos principais?"
- "TL;DR deste documento"

Análise:
- "Quais são os conceitos-chave?"
- "Qual o tema principal?"
- "Explique este documento como se eu tivesse 10 anos"
- "Quais as conclusões do autor?"

Questões:
- "Crie 10 questões de múltipla escolha sobre este conteúdo"
- "Gere 5 questões dissertativas"
- "Crie um quiz com respostas"

Extração:
- "Liste todos os nomes mencionados"
- "Quais datas aparecem no documento?"
- "Extraia todos os números e estatísticas"

🎥 EXEMPLOS DE PERGUNTAS PARA VÍDEOS DO YOUTUBE:

Resumos:
- "Resuma este vídeo"
- "Do que se trata este vídeo?"
- "Principais pontos abordados"
- "Faça um resumo em tópicos"

Análise:
- "Qual a mensagem principal do vídeo?"
- "Quais técnicas/métodos são ensinados?"
- "Liste os exemplos dados no vídeo"

Questões:
- "Crie questões sobre o conteúdo"
- "Teste meu conhecimento sobre o vídeo"

Timestamps:
- "Quais os principais momentos do vídeo?"
- "Crie uma linha do tempo do conteúdo"

💡 DICA: Você pode combinar múltiplas perguntas:
"Resuma o documento, liste os pontos principais e crie 5 questões"
"""

print("[CHAT ANEXOS] ✅ Sistema de anexos e YouTube carregado!")
print("[CHAT ANEXOS] 📄 Suporta: PDF, DOC, DOCX, TXT")
print("[CHAT ANEXOS] 🎥 Suporta: Vídeos do YouTube com legendas")
print("[CHAT ANEXOS] ♾️  SEM LIMITES de tamanho!")
print("[CHAT ANEXOS] 🤖 Usando: Google Gemini API")

# ============================================
# 🚀 INICIALIZAÇÃO
# ============================================
if __name__ == '__main__':
    print("\n" + "=" * 60)
    print("🚀 INICIANDO IAUniev V5.1 FINAL")
    print("=" * 60)

    init_db()

    print("\n📌 SCRAPERS DISPONÍVEIS:")
    if SCRAPER_DISPONIVEL:
        print("   ✅ AVA (Materiais de aula)")
    if LYCEUM_DISPONIVEL:
        print("   ✅ LYCEUM (Notas e faltas)")

    print("\n📌 COMPORTAMENTO:")
    print("   • 1º LOGIN: Scraping automático (5-8 min)")
    print("   • 2º+ LOGIN: Instantâneo (< 5 seg)")
    print("   • BOTÕES: Re-sincroniza quando quiser")
    print("   • SENHA LYCEUM: 9 primeiros dígitos do CPF\n")

    print(f"🌐 Servidor rodando em: http://0.0.0.0:5000")
    print("=" * 60 + "\n")

    app.run(
        debug=True,
        host='0.0.0.0',
        port=5000,
        threaded=True
    )
